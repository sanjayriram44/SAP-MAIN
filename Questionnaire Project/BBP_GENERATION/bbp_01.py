
import os
import re
import docx
import subprocess
import traceback
from dotenv import load_dotenv
from langchain_community.document_loaders import Docx2txtLoader, PyMuPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from core.models import llm_instance, embeddings_instance

# Load env variables
load_dotenv(dotenv_path=".env", override=True)

# Constants
BBP_DIR = "./data"
OUTPUT_DIR = "./output"
QA_PATH = "uploads/qa_input.docx"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "generated_bbp.docx")
CHROMA_DB_DIR = os.path.join(OUTPUT_DIR, "chroma")
PERSONA_PATH = "prompts/persona.txt"
TRIGGER_PATH = "prompts/trigger_prompt.txt"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load BBP documents
bbp_docs = []
for fname in os.listdir(BBP_DIR):
    path = os.path.join(BBP_DIR, fname)
    if fname.endswith(".pdf"):
        bbp_docs.extend(PyMuPDFLoader(path).load())
    elif fname.endswith(".docx"):
        bbp_docs.extend(Docx2txtLoader(path).load())

# Split into chunks
splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
chunks = splitter.split_documents(bbp_docs)

# Load vectorstore with persistence
if os.path.exists(CHROMA_DB_DIR):
    vectorstore = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embeddings_instance)
else:
    vectorstore = Chroma.from_documents(chunks, embeddings_instance, persist_directory=CHROMA_DB_DIR)
    vectorstore.persist()
retriever = vectorstore.as_retriever()

# Load persona + trigger
with open(PERSONA_PATH, "r", encoding="utf-8") as f:
    persona = f.read()
with open(TRIGGER_PATH, "r", encoding="utf-8") as f:
    trigger = f.read()

# Build chat prompt
prompt = ChatPromptTemplate.from_template(f"""
{persona}

{trigger}

Context:
{{context}}

Query:
{{question}}
""")

# Load QA text
qa_text = "\n".join([doc.page_content for doc in Docx2txtLoader(QA_PATH).load()])

# RAG Chain
rag_chain = (
    {"context": retriever | (lambda docs: "\n\n".join([doc.page_content for doc in docs])),
     "question": RunnablePassthrough()}
    | prompt
    | llm_instance
)

try:
    # Invoke chain
    response = rag_chain.invoke(qa_text)
    content = response.content if hasattr(response, 'content') else response

    # Extract Mermaid Diagrams
    mermaid_blocks = re.findall(r"```mermaid\n(.*?)```", content, re.DOTALL)
    image_map = {}
    for idx, code in enumerate(mermaid_blocks):
        mmd = os.path.join(OUTPUT_DIR, f"diagram_{idx+1}.mmd")
        png = os.path.join(OUTPUT_DIR, f"diagram_{idx+1}.png")
        with open(mmd, "w", encoding="utf-8") as f:
            f.write(code.strip())
        try:
            subprocess.run(["mmdc", "-i", mmd, "-o", png, "--puppeteerConfigFile", "puppeteer-config.json"], check=False)
            image_map[f"[[DIAGRAM_{idx+1}]]"] = png
        except Exception:
            print(f"⚠️ Could not render diagram {idx+1}")

    # Replace Mermaid with placeholders
    def repl(match):
        idx = mermaid_blocks.index(match.group(1).strip()) + 1
        return f"[[DIAGRAM_{idx}]]"

    cleaned = re.sub(r"```mermaid\n(.*?)```", repl, content, flags=re.DOTALL)
    cleaned = re.sub(r"[#*`]+", "", cleaned)

    # Save to docx
    doc = docx.Document()
    doc.add_heading("SAP Ariba Sourcing Business Blueprint", 0)
    for para in cleaned.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        elif para.startswith("[[DIAGRAM_"):
            if image_map.get(para):
                doc.add_picture(image_map[para], width=docx.shared.Inches(5.5))
        elif re.match(r"^\d+\. ", para):
            doc.add_heading(para, level=1)
        elif para.startswith("- ") or para.startswith("• "):
            doc.add_paragraph(para.strip("- •"), style="List Bullet")
        else:
            doc.add_paragraph(para)

    doc.save(OUTPUT_PATH)
    print(f"✅ BBP generated and saved to: {OUTPUT_PATH}")

except Exception as e:
    print("❌ BBP generation failed.")
    print(traceback.format_exc())
