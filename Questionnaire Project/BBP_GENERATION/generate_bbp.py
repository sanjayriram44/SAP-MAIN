import os
import re
import docx
import traceback
from dotenv import load_dotenv
from langchain_community.document_loaders import Docx2txtLoader, PyMuPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores.chroma import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from core.models import llm_instance, embeddings_instance
from .renderer import render_mermaid_to_png

# Load environment variables
load_dotenv(dotenv_path=".env", override=True)

# Constants
BBP_DIR = "./BBP_GENERATION/data"
OUTPUT_DIR = "./BBP_GENERATION/output"
QA_PATH = "./BBP_GENERATION/uploads/qa_input.docx"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "generated_bbp.docx")
CHROMA_DB_DIR = os.path.join(OUTPUT_DIR, "chroma")
PERSONA_PATH = "./BBP_GENERATION/prompts/persona.txt"
TRIGGER_PATH = "./BBP_GENERATION/prompts/trigger_prompt.txt"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load BBP reference documents
bbp_docs = []
for fname in os.listdir(BBP_DIR):
    path = os.path.join(BBP_DIR, fname)
    if fname.endswith(".pdf"):
        bbp_docs.extend(PyMuPDFLoader(path).load())
    elif fname.endswith(".docx"):
        bbp_docs.extend(Docx2txtLoader(path).load())

# Split into chunks
splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
chunks = splitter.split_documents(bbp_docs)

# Load or persist Chroma vectorstore
if os.path.exists(CHROMA_DB_DIR):
    vectorstore = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embeddings_instance)
else:
    vectorstore = Chroma.from_documents(chunks, embeddings_instance, persist_directory=CHROMA_DB_DIR)
    vectorstore.persist()

retriever = vectorstore.as_retriever()

# Load prompts
with open(PERSONA_PATH, "r", encoding="utf-8") as f:
    persona = f.read()
with open(TRIGGER_PATH, "r", encoding="utf-8") as f:
    trigger = f.read()

prompt = ChatPromptTemplate.from_template(f"""
{persona}

{trigger}

Context:
{{context}}

Query:
{{question}}
""")

# RAG pipeline
rag_chain = (
    {
        "context": retriever | (lambda docs: "\n\n".join([doc.page_content for doc in docs])),
        "question": RunnablePassthrough()
    }
    | prompt
    | llm_instance
)

# Escape parentheses to prevent Mermaid parser errors
def sanitize_mermaid_code(code):
    return re.sub(r"\((.*?)\)", r"\\(\1\\)", code)

def generate_bbp_from_qa(qa_file_path=QA_PATH):
    try:
        qa_text = "\n".join([doc.page_content for doc in Docx2txtLoader(qa_file_path).load()])
        response = rag_chain.invoke(qa_text)
        content = response.content if hasattr(response, 'content') else response

        # Extract Mermaid diagrams
        mermaid_blocks = re.findall(r"```mermaid\n(.*?)```", content, re.DOTALL)
        image_map = {}

        for idx, code in enumerate(mermaid_blocks):
            sanitized_code = sanitize_mermaid_code(code.strip())
            png_path = render_mermaid_to_png(sanitized_code, OUTPUT_DIR, idx + 1)
            if png_path:
                image_map[f"[[DIAGRAM_{idx+1}]]"] = png_path

        # Replace Mermaid blocks with placeholders
        def repl(match):
            match_str = match.group(1).strip()
            for i, block in enumerate(mermaid_blocks):
                if block.strip() == match_str:
                    return f"[[DIAGRAM_{i+1}]]"
            return "[[DIAGRAM_UNKNOWN]]"

        cleaned = re.sub(r"```mermaid\n(.*?)```", repl, content, flags=re.DOTALL)
        cleaned = re.sub(r"[#*`]+", "", cleaned)

        # Save to Word
        doc = docx.Document()
        doc.add_heading("SAP Ariba Sourcing Business Blueprint", 0)

        for para in cleaned.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            elif para.startswith("[[DIAGRAM_"):
                if image_map.get(para):
                    doc.add_picture(image_map[para], width=docx.shared.Inches(5.5))
            elif re.match(r"^\d+\. ", para):
                doc.add_heading(para, level=1)
            elif para.startswith("- ") or para.startswith("• "):
                doc.add_paragraph(para.strip("- •"), style="List Bullet")
            else:
                doc.add_paragraph(para)

        doc.save(OUTPUT_PATH)
        print(f"✅ BBP generated and saved to: {OUTPUT_PATH}")
        return content, image_map

    except Exception as e:
        print("❌ BBP generation failed.")
        print(traceback.format_exc())
        return None, {}
