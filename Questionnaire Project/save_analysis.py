import os
import docx
from docx.shared import Pt

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def save_individual_and_combined_analysis(pu: str = "", pr: str = ""):
    pu_path = os.path.join(OUTPUT_DIR, "process_understanding.docx")
    pr_path = os.path.join(OUTPUT_DIR, "process_recommendation.docx")
    combined_path = os.path.join(OUTPUT_DIR, "process_analysis_summary.docx")

    # Save PU
    if pu:
        doc_pu = docx.Document()
        doc_pu.add_heading("Process Understanding", level=1)
        doc_pu.add_paragraph(pu).style.font.size = Pt(11)
        doc_pu.save(pu_path)

    # Save PR
    if pr:
        doc_pr = docx.Document()
        doc_pr.add_heading("Process Recommendation", level=1)
        doc_pr.add_paragraph(pr).style.font.size = Pt(11)
        doc_pr.save(pr_path)

    # Save Combined
    if pu or pr:
        doc_combined = docx.Document()
        if pu:
            doc_combined.add_heading("Process Understanding", level=1)
            doc_combined.add_paragraph(pu).style.font.size = Pt(11)
        if pr:
            doc_combined.add_heading("Process Recommendation", level=1)
            doc_combined.add_paragraph(pr).style.font.size = Pt(11)
        doc_combined.save(combined_path)

    print(f"✅ Saved to:\n- {pu_path}\n- {pr_path}\n- {combined_path}")
    return pu_path, pr_path, combined_path
