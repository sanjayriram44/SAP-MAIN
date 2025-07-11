import os
import pandas as pd
import docx

EXCEL_OUTPUT_PATH = "output/conversation_log.xlsx"
WORD_OUTPUT_DIR = "output/qna_docs"

os.makedirs("output", exist_ok=True)
os.makedirs(WORD_OUTPUT_DIR, exist_ok=True)

def save_conversation_to_excel(subprocess_name, conversation_entry):
    """
    Save conversation to an Excel file (one sheet per subprocess).
    """
    data = []
    data.append({"Type": "Main", "Question": conversation_entry["question"], "Answer": conversation_entry["answer"]})

    for fup in conversation_entry.get("followups", []):
        data.append({"Type": "Follow-up", "Question": fup["question"], "Answer": fup["answer"]})

    df_new = pd.DataFrame(data)

    try:
        with pd.ExcelWriter(EXCEL_OUTPUT_PATH, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
            df_new.to_excel(writer, sheet_name=subprocess_name[:31], index=False,
                            header=not writer.sheets.get(subprocess_name[:31]))
    except FileNotFoundError:
        with pd.ExcelWriter(EXCEL_OUTPUT_PATH, engine='openpyxl', mode='w') as writer:
            df_new.to_excel(writer, sheet_name=subprocess_name[:31], index=False)

def save_conversation_to_word(subprocess_name, conversation_entry):
    """
    Save conversation to a Word file per subprocess.
    """
    doc = docx.Document()
    doc.add_heading(f"Q&A for {subprocess_name}", 0)

    # Main Question
    doc.add_heading("Main Question", level=1)
    doc.add_paragraph(conversation_entry["question"])
    doc.add_heading("Answer", level=1)
    doc.add_paragraph(conversation_entry["answer"])

    # Follow-ups
    followups = conversation_entry.get("followups", [])
    if followups:
        doc.add_heading("Follow-Up Questions", level=1)
        for i, fup in enumerate(followups, 1):
            doc.add_paragraph(f"{i}. Q: {fup['question']}", style='List Number')
            doc.add_paragraph(f"   A: {fup['answer']}")

    file_path = os.path.join(WORD_OUTPUT_DIR, f"{subprocess_name[:50]}.docx")
    doc.save(file_path)
